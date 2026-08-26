"""
Step 2a: Unified document extraction.

Dispatches by file type and returns a normalized structure:
    {
        "path": str,
        "doc_type": "pdf" | "html" | "docx",
        "pages": [
            {
                "page_num": int,          # 1-indexed; 0 for non-paginated docs
                "text": str,
                "method": "pymupdf" | "ocr",
                "spans": [ {text, size, bold, bbox} ... ],  # PDF only, used for heading detection
                "tables": [ [[cell, ...], ...], ... ],       # list of tables, each a list of rows
            },
            ...
        ],
        "errors": [str, ...],
    }

PDF extraction strategy:
    1. Try PyMuPDF text extraction per page.
    2. Classify page quality (reusing step1 heuristics). If a page looks
       empty/scanned or garbled, fall back to OCR (pdf2image -> pytesseract)
       for that page only.
    3. Run PyMuPDF's native table finder (page.find_tables()) on every page
       and attach any detected tables.
"""

import re
import statistics
from pathlib import Path

import fitz  # PyMuPDF

MIN_CHARS_PER_PAGE = 40
MIN_ALNUM_RATIO = 0.55
MIN_PRINTABLE_RATIO = 0.85
MAX_OCR_PAGES_PER_DOC = 60  # safety cap so one large scanned book can't stall a batch run

TESSERACT_CMD = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = str(Path(__file__).resolve().parent.parent / "tools" / "poppler_extracted" / "poppler-24.08.0" / "Library" / "bin")


def _classify_text(text: str) -> str:
    stripped = text.strip()
    if not stripped or len(stripped) < MIN_CHARS_PER_PAGE:
        return "EMPTY"
    alnum = sum(1 for c in stripped if c.isalnum())
    printable = sum(1 for c in stripped if c.isprintable() or c in "\n\t")
    alnum_ratio = alnum / len(stripped)
    printable_ratio = printable / len(stripped)
    weird_ratio = len(re.findall(r"[\ufffd\x00-\x08\x0b\x0c\x0e-\x1f]", stripped)) / len(stripped)
    if printable_ratio < MIN_PRINTABLE_RATIO or weird_ratio > 0.02 or alnum_ratio < MIN_ALNUM_RATIO:
        return "GARBLED"
    return "CLEAN"


def _ocr_page(pdf_path: Path, page_num_0indexed: int) -> str:
    import pytesseract
    from pdf2image import convert_from_path

    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
    images = convert_from_path(
        str(pdf_path),
        first_page=page_num_0indexed + 1,
        last_page=page_num_0indexed + 1,
        poppler_path=POPPLER_PATH,
        dpi=200,
    )
    if not images:
        return ""
    return pytesseract.image_to_string(images[0])


def _extract_spans(page) -> list:
    spans = []
    d = page.get_text("dict")
    for block in d.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            line_text = "".join(s["text"] for s in line.get("spans", []))
            if not line_text.strip():
                continue
            first_span = line["spans"][0]
            spans.append({
                "text": line_text,
                "size": round(first_span["size"], 1),
                "bold": bool(first_span["flags"] & 2 ** 4),
                "bbox": line["bbox"],
            })
    return spans


MAX_TABLE_CELL_CHARS = 400  # a "cell" this long is almost certainly a misdetected prose paragraph


def _looks_like_real_table(data: list) -> bool:
    if not data or len(data) < 2:
        return False
    if len(data[0]) < 2:  # single-column "tables" are usually misdetected justified text
        return False
    for row in data:
        for cell in row:
            if cell and len(str(cell)) > MAX_TABLE_CELL_CHARS:
                return False
    return True


def _extract_tables(page) -> list:
    tables = []
    try:
        finder = page.find_tables()
        for tbl in finder.tables:
            data = tbl.extract()
            if _looks_like_real_table(data):
                tables.append(data)
    except Exception:
        pass
    return tables


def extract_pdf(path: Path, max_pages: int | None = None, ocr_on_demand: bool = True) -> dict:
    result = {"path": str(path), "doc_type": "pdf", "pages": [], "errors": []}
    try:
        doc = fitz.open(path)
    except Exception as e:
        result["errors"].append(f"open failed: {type(e).__name__}: {e}")
        return result

    n_pages = doc.page_count
    if max_pages is not None:
        n_pages = min(n_pages, max_pages)

    ocr_pages_used = 0
    for i in range(n_pages):
        try:
            page = doc.load_page(i)
            text = page.get_text("text")
            quality = _classify_text(text)
            method = "pymupdf"

            if quality in ("EMPTY", "GARBLED") and ocr_on_demand:
                if ocr_pages_used >= MAX_OCR_PAGES_PER_DOC:
                    result["errors"].append(
                        f"page {i + 1}: skipped OCR (doc exceeded {MAX_OCR_PAGES_PER_DOC}-page OCR cap)"
                    )
                else:
                    try:
                        ocr_text = _ocr_page(path, i)
                        ocr_pages_used += 1
                        if len(ocr_text.strip()) > len(text.strip()):
                            text = ocr_text
                            method = "ocr"
                    except Exception as e:
                        result["errors"].append(f"page {i + 1} OCR failed: {type(e).__name__}: {e}")

            spans = _extract_spans(page) if method == "pymupdf" else []
            tables = _extract_tables(page) if method == "pymupdf" else []

            result["pages"].append({
                "page_num": i + 1,
                "text": text,
                "method": method,
                "spans": spans,
                "tables": tables,
            })
        except Exception as e:
            result["errors"].append(f"page {i + 1} failed: {type(e).__name__}: {e}")

    doc.close()
    return result


def extract_html(path: Path) -> dict:
    from bs4 import BeautifulSoup

    result = {"path": str(path), "doc_type": "html", "pages": [], "errors": []}
    try:
        raw = path.read_bytes()
        soup = BeautifulSoup(raw, "lxml")
        for tag in soup(["script", "style"]):
            tag.decompose()
        text = soup.get_text("\n")
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        result["pages"].append({"page_num": 0, "text": text, "method": "bs4", "spans": [], "tables": []})
    except Exception as e:
        result["errors"].append(f"{type(e).__name__}: {e}")
    return result


def extract_docx(path: Path) -> dict:
    import docx

    result = {"path": str(path), "doc_type": "docx", "pages": [], "errors": []}
    try:
        d = docx.Document(str(path))
        spans = []
        parts = []
        for para in d.paragraphs:
            if not para.text.strip():
                continue
            parts.append(para.text)
            style = (para.style.name or "").lower()
            if "heading" in style or "title" in style:
                spans.append({"text": para.text, "size": 16.0, "bold": True, "bbox": None})
        text = "\n".join(parts)
        result["pages"].append({"page_num": 0, "text": text, "method": "python-docx", "spans": spans, "tables": []})
    except Exception as e:
        result["errors"].append(f"{type(e).__name__}: {e}")
    return result


def extract_document(path: Path, max_pages: int | None = None) -> dict:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path, max_pages=max_pages)
    if ext in (".htm", ".html"):
        return extract_html(path)
    if ext == ".docx":
        return extract_docx(path)
    return {"path": str(path), "doc_type": ext.lstrip("."), "pages": [], "errors": [f"unsupported extension: {ext}"]}
